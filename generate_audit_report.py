#!/usr/bin/env python3
"""
DATABASE ARCHITECTURE AND FLOW AUDIT REPORT GENERATOR
For the Flowsink project
"""
import os
import sys
from datetime import datetime

# Check for python-docx
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx is not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

OUTPUT_FILE = "DATABASE_ARCHITECTURE_AND_FLOW_REPORT.docx"
PROJECT_NAME = "Flowsink — Student Activity Monitoring & Focus Control Engine"
AUDIT_DATE = datetime.now().strftime("%B %d, %Y")

doc = Document()

# ============================================================
# PAGE SETUP & STYLES
# ============================================================
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Heading styles
for level, size in [(1, 18), (2, 14), (3, 12), (4, 11)]:
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Calibri'
    h_style.font.size = Pt(size)
    h_style.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    h_style.font.bold = True

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def add_paragraph(text, style_name=None, bold=False, italic=False):
    p = doc.add_paragraph("")
    if style_name:
        p.style = style_name
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    run = p.add_run(text)
    return p

def add_code_block(code_text):
    """Add code as shaded paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Add shading (light gray)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F2F2F2"/>')
    pPr.append(shading)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        # Header shading
        tcPr = cell._tc.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="D6E4F0"/>')
        tcPr.append(shading)
    
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    
    doc.add_paragraph()  # spacing
    return table

def add_severity_indicator(level):
    """Add severity color text."""
    colors = {
        'CRITICAL': RGBColor(0xDC, 0x26, 0x26),
        'HIGH': RGBColor(0xF0, 0x7A, 0x00),
        'MEDIUM': RGBColor(0xE6, 0xB8, 0x00),
        'LOW': RGBColor(0x2E, 0x7D, 0x32),
    }
    p = doc.add_paragraph()
    run = p.add_run(f"[{level}]")
    run.bold = True
    run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
    return p

def add_page_break():
    doc.add_page_break()

# ============================================================
# HEADER & FOOTER
# ============================================================
header = section.header
header_para = header.paragraphs[0]
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header_para.add_run("DATABASE ARCHITECTURE AND FLOW AUDIT REPORT")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run(f"Page ")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("DATABASE ARCHITECTURE")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title2.add_run("AND DATA FLOW AUDIT")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

doc.add_paragraph()
doc.add_paragraph()

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_para.add_run(f"Project: {PROJECT_NAME}")
run.font.size = Pt(14)

info_para2 = doc.add_paragraph()
info_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_para2.add_run(f"Audit date: {AUDIT_DATE}")
run.font.size = Pt(12)

info_para3 = doc.add_paragraph()
info_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_para3.add_run("Forensic Read-Only Audit — No Database Modifications Performed")
run.font.size = Pt(11)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()
note_para = doc.add_paragraph()
note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note_para.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
toc_heading = doc.add_heading("TABLE OF CONTENTS", level=1)
toc_items = [
    "1. EXECUTIVE SUMMARY",
    "2. DATABASE INVENTORY",
    "3. SYSTEM ARCHITECTURE",
    "4. DATABASE ARCHITECTURE",
    "5. COMPLETE SCHEMA",
    "6. TABLE-BY-TABLE ANALYSIS",
    "7. RELATIONSHIP / ERD",
    "8. DATA FLOW",
    "9. QUERY FLOW",
    "10. TRANSACTION & CONSISTENCY",
    "11. DATA LIFECYCLE",
    "12. SECURITY AUDIT",
    "13. PERFORMANCE AUDIT",
    "14. LOGIC / CONSISTENCY AUDIT",
    "15. DATABASE DEPENDENCY GRAPH",
    "16. CRITICAL COMPONENTS",
    "17. CRITICAL RISKS",
    "18. RECOMMENDATIONS",
    "19. HOW THE SYSTEM WORKS — FOR NEW DEVELOPERS",
    "20. APPENDIX",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)

add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading("1. EXECUTIVE SUMMARY", level=1)

add_paragraph("This report presents the results of a comprehensive database forensic audit of the Flowsink project — a Student Activity Monitoring & Focus Control Engine. The audit was performed in read-only mode with no modifications made to any data storage, configuration, or production code.", bold=True)

add_paragraph("CRITICAL FINDING:")
p = doc.add_paragraph()
run = p.add_run("The Flowsink project does NOT use any external or persistent database system. ")
run.bold = True
run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
run = p.add_run(
    "All data storage is handled entirely through an in-memory repository "
    "(InMemoryActivityRepository) that exists only for the lifetime of the running process. "
    "No database engine (PostgreSQL, MySQL, SQLite, MongoDB, Redis, etc.) is configured, "
    "connected, or referenced anywhere in the project source code."
)

add_paragraph("Key audit results:")
add_bullet("Databases found: 0 (real persistent databases)")
add_bullet("In-memory data stores identified: 1 (InMemoryActivityRepository)")
add_bullet("Data models/entities: 12 (Pydantic domain models)")
add_bullet("No SQL queries exist anywhere in the project")
add_bullet("No ORM, migration, schema, or seed files found")
add_bullet("No .env or connection-string files found")
add_bullet("No Docker/database configuration files found")
add_bullet("Architecture: Clean Architecture / Hexagonal with in-memory state")

add_paragraph(
    "The persistence design is intentionally minimal: Phase 1 of the project uses an in-memory "
    "repository. The repository interface (ActivityRepository protocol) has been designed so that "
    "SQLite, PostgreSQL, or Redis implementations can be added later without changing the core architecture."
)

add_page_break()

# ============================================================
# 2. DATABASE INVENTORY
# ============================================================
doc.add_heading("2. DATABASE INVENTORY", level=1)

add_paragraph("Comprehensive scan results for all database-related indicators across the entire repository.", bold=True)

add_paragraph("2.1 Database Engines Searched", bold=True)
add_bullet("PostgreSQL — NOT FOUND")
add_bullet("MySQL / MariaDB — NOT FOUND")
add_bullet("SQLite — NOT FOUND (only .gitignore entry excluding future *.sqlite files)")
add_bullet("MongoDB — NOT FOUND")
add_bullet("Redis — NOT FOUND")
add_bullet("Supabase — NOT FOUND")
add_bullet("Firebase — NOT FOUND")
add_bullet("MSSQL / Oracle — NOT FOUND")
add_bullet("Vector database — NOT FOUND")
add_bullet("Cache database — NOT FOUND")
add_bullet("Embedded/local database files (*.db, *.sqlite, *.sqlite3) — NOT FOUND")

add_paragraph("2.2 Connection Indicators Scanned", bold=True)
add_bullet("DATABASE_URL — NOT FOUND")
add_bullet("DB_HOST — NOT FOUND")
add_bullet("DB_PORT — NOT FOUND")
add_bullet("DB_NAME — NOT FOUND")
add_bullet("DB_USER — NOT FOUND")
add_bullet("DB_PASSWORD — NOT FOUND")
add_bullet("Connection strings — NOT FOUND")
add_bullet(".env files — NOT FOUND")
add_bullet("Docker/docker-compose files — NOT FOUND")
add_bullet("SQL files — NOT FOUND")
add_bullet("ORM configurations — NOT FOUND")
add_bullet("Database client libraries — NOT FOUND")
add_bullet("Connection pools — NOT FOUND")
add_bullet("Query builders — NOT FOUND")

add_paragraph("2.3 In-Memory Data Store (The only storage found)", bold=True)
add_table(
    ["Property", "Value"],
    [
        ["Storage type", "In-memory Python lists and dictionaries"],
        ["Class", "InMemoryActivityRepository"],
        ["File", "src/activity_engine/persistence/repository.py"],
        ["Data stored", "ActivityEvent list, StateSnapshot dict, ViolationRecord list"],
        ["Persistence", "NONE — data lost when process exits"],
        ["Max events kept", "1000 (ring buffer)"],
        ["Max violations kept", "1000 (ring buffer)"],
        ["State snapshots", "1 per (student_id, device_id) pair — latest only"],
    ],
    col_widths=[1.8, 4.5]
)

add_paragraph("2.4 Directories and Files Excluded from Scan", bold=True)
add_bullet(".git/ — version control metadata")
add_bullet("node_modules/ — npm dependencies (none installed in this project)")
add_bullet("dist/ / build/ — build artifacts (not present)")
add_bullet("__pycache__/ — Python bytecode cache directories")
add_bullet("Thumbs.db — Windows thumbnail cache (excluded by .gitignore, not present)")

add_page_break()

# ============================================================
# 3. SYSTEM ARCHITECTURE
# ============================================================
doc.add_heading("3. SYSTEM ARCHITECTURE", level=1)

add_paragraph("The system follows Clean Architecture with Hexagonal (Ports & Adapters) principles, plus event-driven internal communication.", bold=True)

add_code_block(
    "+----------------+     +-------------------+     +------------------+\n"
    "|   FRONTEND     |     |    FASTAPI API    |     |  MONITOR ADAPTERS |\n"
    "| (index.html /  | --> |   (server.py)     | --> |  (Process/Browser |\n"
    "|   script.js)   |     | REST + WebSocket  |     |   Network/Mock)   |\n"
    "+----------------+     +-------------------+     +------------------+\n"
    "                              |                            |\n"
    "                              v                            v\n"
    "                     +-------------------+     +------------------+\n"
    "                     |  ACTIVITY ENGINE   |<----|   RAW TELEMETRY  |\n"
    "                     |    (facade.py)     |     |     DICTS        |\n"
    "                     +-------------------+     +------------------+\n"
    "                              |\n"
    "                              v\n"
    "                   +--------------------+\n"
    "                   |   EVENT ENGINE     |\n"
    "                   |   normalize/dedup  |\n"
    "                   +--------------------+\n"
    "                              |\n"
    "                              v\n"
    "                   +--------------------+\n"
    "                   |   POLICY ENGINE    |\n"
    "                   |  evaluate policies |\n"
    "                   +--------------------+\n"
    "                              |\n"
    "                              v\n"
    "                   +--------------------+\n"
    "                   |   STATE ENGINE     |\n"
    "                   |  state machine     |\n"
    "                   +--------------------+\n"
    "                              |\n"
    "                              v\n"
    "                   +--------------------+\n"
    "                   | ESCALATION ENGINE  |\n"
    "                   +--------------------+\n"
    "                              |\n"
    "                              v\n"
    "                   +--------------------+\n"
    "                   |   ACTION ENGINE    |\n"
    "                   +--------------------+\n"
    "                              |\n"
    "                              v\n"
    "                   +--------------------+\n"
    "                   | IN-MEMORY REPOSITORY\n"
    "                   | (NO PERSISTENT DB) |\n"
    "                   +--------------------+"
)

add_paragraph(
    "The frontend (FocusMate) is a standalone HTML/CSS/JS application that calls AI providers "
    "(Gemini/OpenAI-compatible) directly. It is NOT wired to the Python backend. "
    "The Python backend serves as the Activity Engine for student monitoring."
)

add_page_break()

# ============================================================
# 4. DATABASE ARCHITECTURE
# ============================================================
doc.add_heading("4. DATABASE ARCHITECTURE", level=1)

add_paragraph(
    "Since no persistent database exists, the 'database architecture' of this system is the "
    "in-memory repository layer plus the domain model layer that will eventually back a real database."
)

add_paragraph("4.1 Persistence Abstraction (Port)", bold=True)
add_code_block(
    "class ActivityRepository(Protocol):\n"
    "    async def store_event(self, event: ActivityEvent) -> None: ...\n"
    "    async def store_state(self, snapshot: StateSnapshot) -> None: ...\n"
    "    async def store_violation(self, violation: ViolationRecord) -> None: ...\n"
    "    async def get_events(self, limit: int = 100) -> list[ActivityEvent]: ...\n"
    "    async def get_state(self, student_id: str, device_id: str) -> StateSnapshot | None: ...\n"
    "    async def get_violations(self, student_id: str, limit: int = 100) -> list[ViolationRecord]: ..."
)

add_paragraph("4.2 In-Memory Implementation", bold=True)
add_code_block(
    "class InMemoryActivityRepository:\n"
    "    def __init__(self):\n"
    "        self._events: list[ActivityEvent] = []          -- ring buffer max 1000\n"
    "        self._states: dict[tuple[str, str], StateSnapshot] = {}\n"
    "        self._violations: list[ViolationRecord] = []    -- ring buffer max 1000\n"
    "\n"
    "    # Keyed by (student_id, device_id) tuple for states\n"
    "    # Events/violations: append to end; oldest dropped at >1000"
)

add_paragraph("4.3 Planned Future Extensions", bold=True)
add_bullet("SQLite — simple file-based local database (most likely next step)")
add_bullet("PostgreSQL — multi-node production deployment")
add_bullet("Redis — real-time state/event streaming and caching")

add_paragraph("4.4 Storage Layer Boundary", bold=True)
add_table(
    ["Component", "Storage", "Lifetime"],
    [
        ["ActivityRepository protocol", "Interface (abstraction)", "Design-time"],
        ["InMemoryActivityRepository", "RAM (lists/dicts)", "Process lifetime only"],
        ["StateEngine._machines", "RAM (dict)", "Process lifetime only"],
        ["StateEngine._snapshots", "RAM (dict)", "Process lifetime only"],
        ["StateEngine._sessions", "RAM (dict)", "Process lifetime only"],
        ["PolicyEngine._violations", "RAM (list)", "Process lifetime only"],
        ["PolicyEngine._violation_counters", "RAM (dict)", "Process lifetime only"],
        ["Config YAML files", "Disk (read-only)", "Persistent"],
        ["Policy YAML files", "Disk (read-only)", "Persistent"],
    ],
    col_widths=[2.2, 2.2, 1.8]
)

add_page_break()

# ============================================================
# 5. COMPLETE SCHEMA
# ============================================================
doc.add_heading("5. COMPLETE SCHEMA", level=1)

add_paragraph(
    "Since there is no relational database, the 'schema' is defined by Pydantic models. "
    "These models represent the logical data structures that would map to database tables."
)

add_paragraph("5.1 Logical Data Tables (Domain Models)", bold=True)
add_table(
    ["#", "Model", "File", "Purpose"],
    [
        ["1", "ActivityEvent", "core/events.py", "Canonical normalized event flowing through engine"],
        ["2", "ApplicationInfo", "core/events.py", "Application/process details"],
        ["3", "BrowserInfo", "core/events.py", "Browser tab context"],
        ["4", "NetworkInfo", "core/events.py", "DNS/network telemetry"],
        ["5", "StateSnapshot", "core/states.py", "Current state of student/device"],
        ["6", "MonitoringSession", "domain/session.py", "Monitoring session start/end"],
        ["7", "Student", "domain/student.py", "Student identity"],
        ["8", "Device", "domain/device.py", "Device identity"],
        ["9", "DeviceCapabilities", "domain/device.py", "Device feature flags"],
        ["10", "CurrentActivity", "domain/activity.py", "Current activity aggregate"],
        ["11", "ViolationRecord", "domain/violations.py", "Policy violation audit"],
        ["12", "PolicyDecision", "core/decisions.py", "Policy evaluation result"],
        ["13", "ActionRequest", "core/actions.py", "Typed enforcement action request"],
        ["14", "ActionResult", "core/actions.py", "Enforcement action outcome"],
        ["15", "PolicyDocument", "core/policies.py", "Complete policy configuration"],
    ],
    col_widths=[0.4, 1.7, 1.8, 2.3]
)

add_paragraph("5.2 ActivityEvent — Canonical Event Table (Logical)", bold=True)
add_table(
    ["Field", "Type", "Nullable", "Default", "Description"],
    [
        ["event_id", "str", "No", "UUID4", "Unique event identifier"],
        ["device_id", "str", "No", "unknown-device", "Device that produced the event"],
        ["student_id", "str", "Yes", "None", "Student associated with event"],
        ["session_id", "str", "Yes", "None", "Monitoring session identifier"],
        ["timestamp", "datetime", "No", "now UTC", "Event occurrence time"],
        ["source", "EventSource", "No", "-", "process/browser/network/system"],
        ["type", "EventType", "No", "-", "APP_FOCUSED/WEB_NAVIGATION/etc."],
        ["application", "ApplicationInfo", "No", "factory", "Application context"],
        ["browser", "BrowserInfo", "No", "factory", "Browser tab context"],
        ["network", "NetworkInfo", "No", "factory", "Network context"],
        ["metadata", "dict", "No", "{}", "Optional context data"],
        ["schema_version", "int", "No", "1", "Event schema versioning"],
    ],
    col_widths=[1.1, 1.0, 0.6, 1.0, 2.5]
)

add_paragraph("5.3 StateSnapshot — State Table (Logical)", bold=True)
add_table(
    ["Field", "Type", "Nullable", "Default", "Description"],
    [
        ["student_id", "str", "No", "-", "Student identifier"],
        ["device_id", "str", "No", "-", "Device identifier"],
        ["state", "ActivityState", "No", "UNKNOWN", "Current state machine state"],
        ["application", "str", "Yes", "None", "Foreground application"],
        ["domain", "str", "Yes", "None", "Active browser domain"],
        ["started_at", "datetime", "No", "now UTC", "When state began"],
        ["duration_seconds", "float", "No", "0.0", "Duration in current state"],
        ["risk_score", "float", "No", "0.0", "0.0-1.0 risk assessment"],
        ["session_id", "str", "Yes", "None", "Active session"],
        ["updated_at", "datetime", "No", "now UTC", "Last state update"],
    ],
    col_widths=[1.1, 1.0, 0.6, 1.0, 2.5]
)

add_paragraph("5.4 ViolationRecord — Violation Audit Table (Logical)", bold=True)
add_table(
    ["Field", "Type", "Nullable", "Default", "Description"],
    [
        ["violation_id", "str", "No", "UUID4", "Unique violation identifier"],
        ["student_id", "str", "No", "-", "Student who violated policy"],
        ["device_id", "str", "No", "-", "Device used"],
        ["session_id", "str", "Yes", "None", "Monitoring session"],
        ["timestamp", "datetime", "No", "now UTC", "Violation time"],
        ["state_before", "ActivityState", "No", "-", "State before violation"],
        ["state_after", "ActivityState", "No", "-", "State after violation"],
        ["domain", "str", "Yes", "None", "Blocked/flagged domain"],
        ["application", "str", "Yes", "None", "Blocked/flagged app"],
        ["policy_id", "str", "No", "-", "Policy that was violated"],
        ["level", "str", "No", "level_1", "Escalation level"],
        ["action_types", "list[ActionType]", "No", "[]", "Enforcement actions taken"],
        ["reason", "str", "No", "-", "Reason for violation"],
        ["risk_score", "float", "No", "0.0", "Risk score 0.0-1.0"],
    ],
    col_widths=[1.1, 1.2, 0.6, 0.8, 2.5]
)

add_paragraph("5.5 MonitoringSession — Session Table (Logical)", bold=True)
add_table(
    ["Field", "Type", "Nullable", "Default", "Description"],
    [
        ["session_id", "str", "No", "UUID4", "Unique session identifier"],
        ["student_id", "str", "No", "-", "Student in session"],
        ["device_id", "str", "No", "-", "Device being monitored"],
        ["started_at", "datetime", "No", "now UTC", "Session start time"],
        ["ended_at", "datetime", "Yes", "None", "Session end time"],
    ],
    col_widths=[1.1, 1.0, 0.6, 1.2, 2.3]
)

add_page_break()

# ============================================================
# 6. TABLE-BY-TABLE ANALYSIS
# ============================================================
doc.add_heading("6. TABLE-BY-TABLE ANALYSIS", level=1)

add_paragraph("6.1 InMemoryActivityRepository — Central Data Store", bold=True)
add_paragraph(
    "This is the only 'table store' in the system. It holds three collections in RAM: "
    "events (list), states (dict keyed by student+device), and violations (list)."
)
add_table(
    ["Collection", "Type", "Key", "Max Size", "Read Method", "Write Method"],
    [
        ["_events", "list[ActivityEvent]", "append-order", "1000 (ring)", "get_events()", "store_event()"],
        ["_states", "dict[(student,dev)]", "composite", "1 per key", "get_state()", "store_state()"],
        ["_violations", "list[ViolationRecord]", "append-order", "1000 (ring)", "get_violations()", "store_violation()"],
    ],
    col_widths=[1.3, 1.7, 1.2, 1.0, 1.0, 1.0]
)

add_paragraph("6.2 StateEngine Internal Stores", bold=True)
add_table(
    ["Store", "Type", "Key", "Purpose"],
    [
        ["_machines", "dict[(student,dev)] → ActivityStateMachine", "composite", "State machine per student/device"],
        ["_snapshots", "dict[(student,dev)] → StateSnapshot", "composite", "Latest state snapshot"],
        ["_sessions", "dict[(student,dev)] → MonitoringSession", "composite", "Active monitoring session"],
    ],
    col_widths=[1.4, 2.6, 1.2, 2.0]
)

add_paragraph("6.3 PolicyEngine Internal Stores", bold=True)
add_table(
    ["Store", "Type", "Key", "Purpose"],
    [
        ["_violation_counters", "dict[(student,dev)] → int", "composite", "Running violation count for escalation"],
        ["_activity_started_at", "dict[(student,dev)] → datetime", "composite", "When current violation streak began"],
        ["_violations", "list[ViolationRecord]", "append-order", "Violation audit trail (max 1000)"],
    ],
    col_widths=[1.7, 2.5, 1.2, 1.8]
)

add_paragraph("6.4 EventEngine Internal Store", bold=True)
add_table(
    ["Store", "Type", "Key", "Purpose"],
    [
        ["_last_seen", "dict[str, float]", "dedupe_key", "Timestamp tracking for debounce dedup (max 4096)"],
    ],
    col_widths=[1.2, 1.5, 1.2, 3.3]
)

add_paragraph("6.5 ActionEngine Internal Store", bold=True)
add_table(
    ["Store", "Type", "Key", "Purpose"],
    [
        ["_completed", "set[(action, target, student)]", "composite", "Idempotency dedup for executed actions"],
        ["_restricted_mode_active", "bool", "-", "Tracks restricted mode state"],
    ],
    col_widths=[1.5, 1.8, 1.2, 2.7]
)

add_paragraph("6.6 ActivityService Internal Store", bold=True)
add_table(
    ["Store", "Type", "Key", "Purpose"],
    [
        ["_current", "CurrentActivity | None", "-", "Latest computed current activity aggregate"],
    ],
    col_widths=[1.2, 2.0, 1.0, 3.0]
)

add_paragraph("6.7 Frontend localStorage / In-Memory State", bold=True)
add_table(
    ["Store", "Type", "Purpose"],
    [
        ["localStorage gemini_api_key", "string", "Stores Gemini API key on user's browser"],
        ["timetableData", "JS array (RAM)", "Weekly schedule grid in browser memory"],
        ["timeSlots", "JS array (RAM)", "Custom time slot definitions"],
        ["customActivities", "JS array (RAM)", "User custom activity types"],
        ["analysisResult", "JS object (RAM)", "Last AI analysis result"],
    ],
    col_widths=[2.0, 1.6, 3.6]
)

add_page_break()

# ============================================================
# 7. RELATIONSHIP / ERD
# ============================================================
doc.add_heading("7. RELATIONSHIP / ERD", level=1)

add_paragraph(
    "Since there is no SQL database, the relationships below are logical relationships "
    "between Pydantic models and in-memory stores."
)

add_code_block(
    "                     LOGICAL RELATIONSHIP DIAGRAM\n"
    "                     ============================\n"
    "\n"
    "  STUDENT (domain/student.py)\n"
    "    |\n"
    "    | 1:N  (one student, multiple monitoring sessions)\n"
    "    v\n"
    "  MONITORING_SESSION (domain/session.py)\n"
    "    |\n"
    "    | N:1  (many sessions reference one device)\n"
    "    v\n"
    "  DEVICE (domain/device.py)\n"
    "    |\n"
    "    | N:1  (many state snapshots reference one device)\n"
    "    v\n"
    "  STATE_SNAPSHOT (core/states.py) -- unique key: (student_id, device_id)\n"
    "    |\n"
    "    | N:1  (one event can update one snapshot)\n"
    "    v\n"
    "  ACTIVITY_EVENT (core/events.py)\n"
    "    |\n"
    "    | N:1  (each event may generate one policy decision)\n"
    "    v\n"
    "  POLICY_DECISION (core/decisions.py)\n"
    "    |\n"
    "    | N:N  (one decision can produce multiple actions)\n"
    "    v\n"
    "  ACTION_REQUEST (core/actions.py)\n"
    "    |\n"
    "    | N:1  (each action request yields one result)\n"
    "    v\n"
    "  ACTION_RESULT (core/actions.py)\n"
    "\n"
    "  VIOLATION_RECORD (domain/violations.py)\n"
    "    |\n"
    "    | N:1 (references student, device, session)\n"
    "    +---> references ActivityState (state_before, state_after)\n"
    "    +---> references PolicyDecision data (level, actions, reason)\n"
)

add_paragraph("Key Relationships:", bold=True)
add_bullet("Student 1:N MonitoringSession (a student may have many sessions over time)")
add_bullet("Device 1:N MonitoringSession (a device may host many sessions)")
add_bullet("(student_id, device_id) is the composite key for StateSnapshot, state machine, sessions")
add_bullet("ActivityEvent N:1 PolicyDecision (each event is evaluated once)")  
add_bullet("PolicyDecision N:N ActionRequest (a decision can plan multiple actions)")
add_bullet("ActionRequest 1:1 ActionResult (each request produces exactly one result)")
add_bullet("ViolationRecord references Student, Device, Policy; 1:N per student")

add_page_break()

# ============================================================
# 8. DATA FLOW
# ============================================================
doc.add_heading("8. DATA FLOW", level=1)

add_paragraph("8.1 Complete End-to-End Data Flow", bold=True)

add_code_block(
    "TELEMETRY INPUT FLOW:\n"
    "=====================\n"
    "Monitor Adapters (Process/Browser/Network/Mock)\n"
    "    |\n"
    "    |  raw dict e.g. {\"kind\": \"browser_navigation\", \"browser\": {\"domain\": \"youtube.com\"}}\n"
    "    v\n"
    "MonitoringService (services/monitoring_service.py)\n"
    "    |\n"
    "    v\n"
    "EventEngine.process_raw (engine/event_engine.py)\n"
    "    |  validate -> normalize -> dedupe -> drop-stale\n"
    "    v\n"
    "ActivityEvent (canonical)\n"
    "    |\n"
    "    v\n"
    "Facade._route subscriber\n"
    "    |\n"
    "    v\n"
    "PolicyEngine.evaluate (engine/policy_engine.py)\n"
    "    |  evaluate against PolicyDocument\n"
    "    |  escalate if violation count/duration triggers\n"
    "    v\n"
    "PolicyDecision\n"
    "    |\n"
    "    +---> StateEngine.apply_decision\n"
    "    |        |  transition ActivityStateMachine\n"
    "    |        v\n"
    "    |      StateSnapshot\n"
    "    |        |\n"
    "    |        v\n"
    "    |      InMemoryActivityRepository.store_state\n"
    "    |\n"
    "    +---> EscalationEngine.plan\n"
    "    |        |\n"
    "    |        v\n"
    "    |      List[ActionRequest]\n"
    "    |        |\n"
    "    |        v\n"
    "    |      ActionEngine.execute\n"
    "    |        |  respects mode (dry_run/audit/enforce)\n"
    "    |        |  checks idempotency\n"
    "    |        v\n"
    "    |      ActionResult (returned to adapter)\n"
    "    |\n"
    "    v\n"
    "InMemoryActivityRepository.store_event\n"
    "\n"
    "REQUERY FLOW:\n"
    "=============\n"
    "API GET /api/events\n"
    "    -> ActivityEngine.recent_events\n"
    "    -> InMemoryActivityRepository.get_events\n"
    "    -> list[ActivityEvent]\n"
    "\n"
    "API GET /api/state\n"
    "    -> ActivityEngine.current_state\n"
    "    -> StateEngine.get_snapshot\n"
    "    -> StateSnapshot | None\n"
    "\n"
    "API GET /api/metrics\n"
    "    -> ActivityEngine.metrics\n"
    "    -> dict from all engines\n"
)

add_paragraph("8.2 API Interface — REST Endpoints", bold=True)
add_table(
    ["Method", "Endpoint", "Controller", "Service", "Data Source", "Output"],
    [
        ["GET", "/", "health()", "-", "-", "Status JSON"],
        ["POST", "/api/session/start", "start_session()", "ActivityEngine.start_session()", "StateEngine._sessions", "session_id + started_at"],
        ["POST", "/api/session/end", "end_session()", "ActivityEngine.end_session()", "StateEngine._sessions", "status + session_id"],
        ["GET", "/api/state", "get_state()", "ActivityEngine.current_state()", "StateEngine._snapshots", "StateSnapshot JSON"],
        ["GET", "/api/metrics", "get_metrics()", "ActivityEngine.metrics()", "Engine metrics dicts", "Metrics JSON"],
        ["GET", "/api/events", "get_events()", "ActivityEngine.recent_events()", "InMemoryActivityRepository._events", "Event JSON list"],
        ["POST", "/api/telemetry", "feed_telemetry()", "ActivityEngine.feed_raw()", "-> EventEngine-> PolicyEngine-> StateEngine", "Status + event_id + state"],
        ["WS", "/ws", "websocket_endpoint()", "-", "All broadcasts", "HELLO/EVENT/STATE/PONG messages"],
    ],
    col_widths=[0.6, 1.6, 1.4, 1.6, 1.5, 1.5]
)

add_paragraph("8.3 Frontend — Backend Disconnect", bold=True)
add_paragraph(
    "The frontend (index.html, script.js) is a standalone application with NO connection to the Python backend. "
    "It directly calls Gemini (Google) or OpenAI-compatible REST APIs from the browser. "
    "Data flows entirely in-browser: user inputs schedule -> JS builds prompt -> AI API called -> result rendered in DOM."
)

add_page_break()

# ============================================================
# 9. QUERY FLOW
# ============================================================
doc.add_heading("9. QUERY FLOW", level=1)

add_paragraph(
    "No SQL queries exist in this project. The 'queries' are repository method calls "
    "and engine evaluation functions. Below is the source-code to data-store mapping."
)

add_paragraph("9.1 Read Operations", bold=True)
add_table(
    ["Source File", "Function/Method", "Operation", "Data Store", "Type"],
    [
        ["persistence/repository.py", "get_events()", "list(reversed(self._events[-limit:]))", "_events", "READ"],
        ["persistence/repository.py", "get_state()", "dict lookup by (student_id, device_id)", "_states", "READ"],
        ["persistence/repository.py", "get_violations()", "filter by student_id, reversed, slice", "_violations", "READ"],
        ["engine/state_engine.py", "get_snapshot()", "dict lookup by (student_id, device_id)", "_snapshots", "READ"],
        ["engine/state_engine.py", "get_session()", "dict lookup by (student_id, device_id)", "_sessions", "READ"],
        ["engine/policy_engine.py", "violations_for()", "filter reversed _violations by student_id", "_violations", "READ"],
        ["engine/facade.py", "recent_events()", "await repository.get_events(limit)", "repository", "READ"],
        ["engine/facade.py", "violations()", "await repository.get_violations()", "repository", "READ"],
        ["engine/facade.py", "current_state()", "state_engine.get_snapshot()", "state_engine", "READ"],
        ["engine/facade.py", "metrics()", "dict of engine metrics", "engines", "READ"],
    ],
    col_widths=[1.8, 1.8, 2.3, 1.2, 0.6]
)

add_paragraph("9.2 Write Operations", bold=True)
add_table(
    ["Source File", "Function/Method", "Operation", "Data Store", "Type"],
    [
        ["persistence/repository.py", "store_event()", "append + trim >1000", "_events", "WRITE"],
        ["persistence/repository.py", "store_state()", "dict[key] = snapshot", "_states", "UPSERT"],
        ["persistence/repository.py", "store_violation()", "append + trim >1000", "_violations", "WRITE"],
        ["engine/state_engine.py", "apply_decision()", "transition + set snapshot", "_snapshots", "UPSERT"],
        ["engine/state_engine.py", "force_state()", "reset + transition + set snapshot", "_snapshots", "UPSERT"],
        ["engine/state_engine.py", "start_session()", "dict[key] = MonitoringSession", "_sessions", "WRITE"],
        ["engine/state_engine.py", "end_session()", "update + reset machine + delete session", "_sessions", "UPDATE"],
        ["engine/policy_engine.py", "evaluate()", "increment violation counters", "_violation_counters", "UPDATE"],
        ["engine/policy_engine.py", "_record_violation()", "append + trim >1000", "_violations", "WRITE"],
        ["engine/action_engine.py", "execute()", "add to _completed set", "_completed", "WRITE"],
        ["engine/facade.py", "_route()", "store_event + store_state", "repository", "WRITE"],
        ["engine/event_engine.py", "_is_duplicate()", "update _last_seen dict", "_last_seen", "UPSERT"],
    ],
    col_widths=[1.8, 1.8, 2.5, 1.4, 0.6]
)

add_paragraph("9.3 Delete Operations", bold=True)
add_table(
    ["Source File", "Function/Method", "Operation", "Data Store", "Type"],
    [
        ["engine/event_engine.py", "_is_duplicate()", "prune keys older than 1 hour if >4096", "_last_seen", "DELETE (GC)"],
        ["engine/state_engine.py", "end_session()", "machine.reset()", "_machines", "RESET"],
        ["engine/policy_engine.py", "reset_violations()", "clear counters + activity_started_at", "_violation_counters", "DELETE"],
        ["persistence/repository.py", "store_event()", "drop oldest when >1000", "_events", "DELETE (ring)"],
        ["persistence/repository.py", "store_violation()", "drop oldest when >1000", "_violations", "DELETE (ring)"],
    ],
    col_widths=[1.8, 1.8, 2.8, 1.3, 0.9]
)

add_page_break()

# ============================================================
# 10. TRANSACTION & CONSISTENCY
# ============================================================
doc.add_heading("10. TRANSACTION & CONSISTENCY", level=1)

add_paragraph(
    "No transactional database operations exist because there is no database. "
    "However, the in-memory operations have consistency properties worth analyzing."
)

add_paragraph("10.1 Atomicity Analysis", bold=True)
add_table(
    ["Operation", "Atomic?", "Notes"],
    [
        ["store_event()", "Yes", "Single list append"],
        ["store_state()", "Yes", "Single dict assignment"],
        ["store_violation()", "Yes", "Single list append"],
        ["apply_decision()", "No", "State transition + snapshot creation are two steps but no observers can interleave in sync code"],
        ["_route() in facade", "No", "Calls store_event, then store_state, then executes actions. If state store fails, event is already stored."],
        ["ActionEngine.execute()", "No", "Multiple state updates (status, metrics, _completed set)"],
    ],
    col_widths=[2.0, 1.0, 3.8]
)

add_paragraph("10.2 Consistency Issues Found", bold=True)

p = doc.add_paragraph()
run = p.add_run("MEDIUM — Event stored before state in facade._route()")
run.bold = True
add_bullet("Location: engine/facade.py lines 86-98")
add_bullet("Issue: In _route(), store_event() is awaited before store_state(). If store_state() fails, event exists without corresponding state.")
add_bullet("Risk: Inconsistent snapshot between events and states")

p = doc.add_paragraph()
run = p.add_run("MEDIUM — Violations tracked in two locations")
run.bold = True
add_bullet("Location: PolicyEngine._violations AND InMemoryActivityRepository._violations")
add_bullet("Issue: PolicyEngine maintains its own _violations list but repository.store_violation() is never called by PolicyEngine.")
add_bullet("Evidence: repository.store_violation exists as protocol method but no caller invokes it except potential future callers.")
add_bullet("Impact: The repository's violation storage is effectively dead code; violations are only accessible via PolicyEngine.violations_for().")

p = doc.add_paragraph()
run = p.add_run("LOW — State snapshots can be overwritten")
run.bold = True
add_bullet("Location: StateEngine._snapshots dict uses (student_id, device_id) as key")
add_bullet("Issue: Only one snapshot per (student, device) retained — no history.")
add_bullet("Impact: Past state history is lost unless events are replayed.")

p = doc.add_paragraph()
run = p.add_run("LOW — Monitoring sessions dictionary doesn't remove ended sessions")
run.bold = True
add_bullet("Location: StateEngine.end_session() updates the session in place")
add_bullet("Issue: Ended sessions remain in _sessions dict forever (no cleanup)")
add_bullet("Impact: Memory grows with number of sessions during long-running processes.")

add_paragraph("10.3 Race Condition Analysis", bold=True)
add_bullet("The system is asyncio-based and relies on single-threaded event loop execution.")
add_bullet("The _route() subscriber callback in facade.py is async; multiple events may interleave between awaits.")
add_bullet("store_event() and store_state() are not wrapped in a lock. Between the two awaits, another event could store different state.")
add_bullet("StateEngine.apply_decision() performs synchronous operations — atomic within single coroutine execution (no await), but the snapshot stored by repository may differ from the one computed by state_engine if another coroutine interleaves between the state_engine.apply_decision() and repository.store_state() calls.")

add_paragraph("10.4 Idempotency", bold=True)
add_table(
    ["Component", "Idempotent?", "Mechanism"],
    [
        ["EventEngine", "Yes", "dedupe_key + debounce window"],
        ["ActionEngine", "Yes", "_completed set with (action, target, student)"],
        ["StateEngine transitions", "Yes", "Target == current state returns same state"],
        ["PolicyEngine escalation", "Partial", "Violation counters can be reset by ALLOWED/FOCUS outcomes"],
    ],
    col_widths=[2.2, 1.2, 3.4]
)

add_page_break()

# ============================================================
# 11. DATA LIFECYCLE
# ============================================================
doc.add_heading("11. DATA LIFECYCLE", level=1)

add_paragraph("11.1 Data Creation", bold=True)
add_table(
    ["Data Type", "Created By", "Creation Path", "File"],
    [
        ["ActivityEvent", "EventEngine.normalize()", "Monitor -> process_raw -> normalize", "engine/event_engine.py"],
        ["StateSnapshot", "StateEngine.apply_decision()", "PolicyDecision -> transition -> snapshot", "engine/state_engine.py"],
        ["ViolationRecord", "PolicyEngine._record_violation()", "WARNING outcome -> record", "engine/policy_engine.py"],
        ["MonitoringSession", "StateEngine.start_session()", "API POST /api/session/start", "engine/state_engine.py"],
        ["PolicyDecision", "PolicyEvaluator.evaluate()", "ActivityEvent -> policy match", "policy/evaluator.py"],
        ["ActionRequest", "EscalationEngine.plan()", "PolicyDecision action_types -> requests", "engine/escalation_engine.py"],
        ["ActionResult", "ActionEngine.execute()", "Adapter method -> result", "engine/action_engine.py"],
        ["CurrentActivity", "ActivityService.update_from_event()", "ActivityEvent -> aggregate", "services/activity_service.py"],
    ],
    col_widths=[1.2, 1.5, 2.5, 1.6]
)

add_paragraph("11.2 Data Validation", bold=True)
add_bullet("Raw events: validated via pydantic ActivityEvent construction in EventEngine.normalize()")
add_bullet("Policy YAML: validated via PolicyDocument.model_validate() in PolicyLoader.from_dict()")
add_bullet("Config YAML: validated via Config.model_validate() in ConfigLoader.load()")
add_bullet("Action requests: validated via pydantic ActionRequest/typed payload models")
add_bullet("Domain models: all frozen pydantic BaseModel with type validation")

add_paragraph("11.3 Data Storage", bold=True)
add_bullet("All runtime data stored in RAM (in-memory lists/dicts)")
add_bullet("No persistent storage — data lost on process termination/restart")
add_bullet("Frontend data stored in browser memory/localStorage")

add_paragraph("11.4 Data Read", bold=True)
add_table(
    ["Data Consumer", "Reads", "Purpose"],
    [
        ["API /api/events", "ActivityEvent list", "Display recent events"],
        ["API /api/state", "StateSnapshot", "Display current student state"],
        ["API /api/metrics", "Engine metrics dicts", "Display engine health"],
        ["API /api/session/start", "MonitoringSession", "Start monitoring"],
        ["API /api/session/end", "MonitoringSession", "End monitoring"],
        ["FocusService", "StateSnapshot", "Get current focus state"],
        ["BedtimeService", "StateEngine.force_state()", "Apply bedtime states"],
        ["ActivityService", "CurrentActivity", "Track current activity"],
    ],
    col_widths=[2.0, 1.8, 2.6]
)

add_paragraph("11.5 Data Update", bold=True)
add_bullet("StateSnapshot: overwritten on every policy decision via apply_decision()")
add_bullet("MonitoringSession: updated only once at end_session()")
add_bullet("Violation counters: incremented/decremented on each evaluation")
add_bullet("stored events/violations: append-only (never updated in place)")

add_paragraph("11.6 Data Deletion / Expiry", bold=True)
add_bullet("Events > 1000: oldest dropped when new events appended (ring buffer)")
add_bullet("Violations > 1000: oldest dropped when new violations appended (ring buffer)")
add_bullet("EventEngine._last_seen > 4096: prune keys older than 1 hour")
add_bullet("State snapshots: overwritten (no history)")
add_bullet("Sessions: never deleted from _sessions dict (only reset via end_session)")
add_bullet("No soft delete, retention policy, archive mechanism, or backup system exists")

add_paragraph("11.7 Backup / Recovery", bold=True)
add_bullet("NO backup mechanism exists — data is ephemeral by design")
add_bullet("NO recovery mechanism — restarting the process starts with empty state")
add_bullet("NO audit logs written to disk — only structured logging via Python logging module")

add_page_break()

# ============================================================
# 12. SECURITY AUDIT
# ============================================================
doc.add_heading("12. SECURITY AUDIT", level=1)

add_paragraph(
    "No database credentials or secrets were found because no database connection exists. "
    "However, several security-relevant observations were identified."
)

add_paragraph("12.1 Credentials / Secrets", bold=True)
add_table(
    ["Item", "Status", "Details"],
    [
        ["DATABASE_URL", "NOT FOUND", "No database connection string"],
        ["DB credentials", "NOT FOUND", "No database credentials"],
        [".env files", "NOT FOUND", "No .env, .env.local, or .env.* files in repository"],
        ["API keys in code", "NONE", "No hard-coded API keys found"],
        ["Frontend API key loading", "BROWSER", "script.js fetches .env.local or localStorage for Gemini key"],
        ["AI provider keys", "USER-PROVIDED", "User enters Gemini/OpenAI keys in browser (stored in localStorage or .env.local)"],
    ],
    col_widths=[2.0, 1.4, 3.4]
)

add_paragraph("12.2 Security Issues Found", bold=True)

p = doc.add_paragraph()
run = p.add_run("HIGH — Sensitive API key stored in localStorage")
run.bold = True
run.font.color.rgb = RGBColor(0xF0, 0x7A, 0x00)
add_bullet("File: script.js, getGeminiKey() function, line 463")
add_bullet("Problem: Gemini API key is stored in browser localStorage when input by user")
add_bullet("Risk: XSS attack could exfiltrate the key; any browser extension/user with access to the browser profile can read it")
add_bullet("Recommendation: Use a backend proxy to proxy AI requests; do not expose API keys to frontend")

p = doc.add_paragraph()
run = p.add_run("MEDIUM — .env.local served as static file")
run.bold = True
run.font.color.rgb = RGBColor(0xE6, 0xB8, 0x00)
add_bullet("File: index.html + script.js (getGeminiKey)")
add_bullet("Problem: script.js fetches '.env.local' from the web root. If deployed as static files, this file is publicly accessible.")
add_bullet("Risk: API key leakage to anyone who can access the site")
add_bullet("Recommendation: Never serve .env.local publicly; use a server-side proxy")

p = doc.add_paragraph()
run = p.add_run("MEDIUM — Python FastAPI has no authentication")
run.bold = True
run.font.color.rgb = RGBColor(0xE6, 0xB8, 0x00)
add_bullet("File: server.py")
add_bullet("Problem: All REST and WebSocket endpoints are public, no auth required")
add_bullet("Risk: Unauthorized users can start/stop sessions, inject telemetry, read state")
add_bullet("Recommendation: Add API key / bearer token authentication")

p = doc.add_paragraph()
run = p.add_run("LOW — CORS is wide open for methods/headers")
run.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
add_bullet("File: server.py lines 104-110")
add_bullet("Problem: allow_methods=['*'], allow_headers=['*']")
add_bullet("Risk: Broad CORS policy could allow cross-origin requests")
add_bullet("Recommendation: Restrict methods/headers to what's actually used")

add_paragraph("12.3 Sensitive Data Handling", bold=True)
add_table(
    ["Data", "Sensitivity", "Handling"],
    [
        ["API Keys", "HIGH", "Stored in localStorage, fetched from .env.local — insecure"],
        ["Student IDs", "MEDIUM", "Default values like 'student-001'; no PII"],
        ["Browser domains/URLs", "MEDIUM", "Stored in events; domain-only by default; URL only if policy allows"],
        ["Window titles", "LOW", "Captured in ApplicationInfo (privacy risk if not filtered)"],
        ["Device IDs", "LOW", "Auto-generated or from hostname"],
        ["Session IDs", "LOW", "UUID4 generated"],
    ],
    col_widths=[1.8, 1.2, 3.8]
)

add_paragraph("12.4 SQL Injection / Raw SQL", bold=True)
add_bullet("NO SQL injection risk: no SQL exists in the codebase")
add_bullet("No raw SQL queries, no string-built queries, no SQL anywhere")

add_page_break()

# ============================================================
# 13. PERFORMANCE AUDIT
# ============================================================
doc.add_heading("13. PERFORMANCE AUDIT", level=1)

add_paragraph("Performance analysis of the in-memory data structures and their operations.")

add_paragraph("13.1 Complexity Analysis", bold=True)
add_table(
    ["Operation", "Time Complexity", "Space Complexity", "Notes"],
    [
        ["store_event()", "O(1) append", "O(n) up to 1000", "Ring buffer prevents unbounded growth"],
        ["store_state()", "O(1) dict set", "O(1) per key", "1 per (student, device)"],
        ["store_violation()", "O(1) append", "O(n) up to 1000", "Ring buffer"],
        ["get_events(limit)", "O(n) slice+reverse", "O(limit)", "Copies list on read"],
        ["get_state()", "O(1) dict get", "O(1)", "Efficient"],
        ["get_violations()", "O(n) filter", "O(limit)", "Linear scan of violations list"],
        ["StateEngine.apply_decision()", "O(1)", "O(1)", "Dict operations"],
        ["EventEngine._is_duplicate()", "O(1)", "O(4096) max", "Debounce cache with pruning"],
        ["PolicyEngine.evaluate()", "O(p) policy size", "O(1)", "Domain/app pattern matching"],
        ["ActionEngine.execute()", "O(1)", "O(1)", "Idempotency set check"],
    ],
    col_widths=[2.2, 1.6, 1.6, 2.2]
)

add_paragraph("13.2 Performance Issues", bold=True)

p = doc.add_paragraph()
run = p.add_run("LOW — get_violations() is O(n) linear scan")
run.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
add_bullet("File: persistence/repository.py, line 53-54")
add_bullet("Every call scans the full violations list to filter by student_id")
add_bullet("Max 1000 items — acceptable for current scale")

p = doc.add_paragraph()
run = p.add_run("LOW — get_events() copies list every call")
run.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
add_bullet("File: persistence/repository.py, line 48")
add_bullet("list(reversed(self._events[-limit:])) creates a copy")
add_bullet("Max 1000 items — negligible")

p = doc.add_paragraph()
run = p.add_run("MEDIUM — No pagination for event/violation reads")
run.bold = True
run.font.color.rgb = RGBColor(0xE6, 0xB8, 0x00)
add_bullet("API GET /api/events supports limit but no offset/cursor")
add_bullet("Frontend cannot page through old events")

p = doc.add_paragraph()
run = p.add_run("MEDIUM — Full event history lost after 1000")
run.bold = True
run.font.color.rgb = RGBColor(0xE6, 0xB8, 0x00)
add_bullet("Ring buffer drops oldest events after 1000")
add_bullet("No audit trail for activity beyond the most recent 1000 events")

add_paragraph("13.3 N+1 Query Risk", bold=True)
add_bullet("NOT APPLICABLE — No SQL/ORM queries exist")
add_bullet("The in-memory operations are all O(1) or O(n) single-pass")

add_page_break()

# ============================================================
# 14. LOGIC / CONSISTENCY AUDIT
# ============================================================
doc.add_heading("14. LOGIC / CONSISTENCY AUDIT", level=1)

add_paragraph("Checking for mapping consistency between models, code, and schema.")

add_paragraph("14.1 Model vs Code Mapping", bold=True)
add_table(
    ["COMPONENT", "EXPECTED", "ACTUAL", "MISMATCH", "SEVERITY"],
    [
        ["ActivityRepository protocol", "store_violation() defined", "method exists in protocol", "NONE", "-"],
        ["InMemoryActivityRepository", "implements all protocol methods", "implements store_event/store_state/store_violation/get_events/get_state/get_violations", "NONE — full implementation", "-"],
        ["Facade calls store_violation", "Should call repository.store_violation()", "Never called — PolicyEngine has own _violations list", "YES — dead interface method", "MEDIUM"],
        ["StateSnapshot used by API", "/api/state returns snapshot.model_dump()", "server.py uses model_dump()", "NONE", "-"],
        ["ActivityEngine.metrics()", "Contains event_engine + policy_engine + action_engine", "Returns all three", "NONE", "-"],
        ["MonitoringSession.end()", "Returns session with ended_at", "model_copy update ended_at", "NONE", "-"],
        ["FocusService requires 5 engines", "Has event/policy/state/action/escalation", "All present", "NONE", "-"],
        ["BedtimeService._issue_warning", "Creates ActionRequest with payload", "Creates dynamic type 'WarningRequest' instead of real WarningRequest class", "TYPE MISMATCH — uses type() instead of WarningRequest", "LOW"],
        ["EventEngine.process() skips dedupe on duplicates", "process() should work like process_raw()", "process() increments events_duplicate but still notifies subscribers", "INCONSISTENT — process() notifies on duplicates", "MEDIUM"],
        ["PolicyEngine.violations_for()", "Uses same source as repository", "Reads from its own _violations list", "DUPLICATE STORAGE", "MEDIUM"],
        ["ActionEngine._completed set never clears", "Set grows unbounded", "No pruning mechanism", "MEMORY LEAK over time", "LOW"],
        ["StateEngine._sessions never prunes", "Ended sessions stay in dict", "No cleanup", "MEMORY LEAK over time", "LOW"],
        ["EventEngine._last_seen prunes", "Keys pruned when >4096", "Prunes by cutoff 1 hour", "OK", "-"],
    ],
    col_widths=[1.7, 1.5, 2.2, 1.5, 0.9]
)

add_paragraph("14.2 Dead Code / Unused Elements", bold=True)
add_table(
    ["Item", "Location", "Status"],
    [
        ["ActivityRepository.store_violation()", "persistence/repository.py", "Never called"],
        ["ActivityRepository.get_violations()", "persistence/repository.py", "Never called from facade"],
        ["ViolationRecord state_after", "domain/violations.py", "Always set same as state_before in _record_violation()"],
        ["DecisionMessage / ActionMessage", "transport/dto.py", "Defined but never used in server.py broadcasts"],
        ["DeviceController port", "ports/__init__.py", "Interface defined; only mocked, never used by core"],
        ["EventSink port", "ports/__init__.py", "Defined but never implemented"],
        ["Clock port", "ports/__init__.py", "Defined but never implemented"],
        ["MockNetworkMonitor._domains", "adapters/mock/network_monitor.py", "Set only via add_domain; add_domain never used in tests/sim"],
        ["allowed_url_patterns", "core/policies.py", "Defined in FocusPolicy; never used in evaluator"],
        ["ProcessMonitor.get_running_processes", "ports/__init__.py", "Defined but never used in monitoring service"],
    ],
    col_widths=[2.4, 2.2, 2.2]
)

add_paragraph("14.3 Naming Consistency", bold=True)
add_table(
    ["Pattern", "Status", "Notes"],
    [
        ["snake_case Python", "CONSISTENT", "All Python follows PEP8"],
        ["camelCase JS", "CONSISTENT", "All JS follows camelCase"],
        ["Enum values uppercase", "CONSISTENT", "EventType, ActionType, ActivityState use ALL_CAPS"],
        ["Teacher/student identifiers", "CONSISTENT", "student_id used throughout"],
        ["device identification", "PARTIALLY CONSISTENT", "device_id used; some places use 'device-001', some use 'unknown-device', some use hostname"],
        ["Session identifiers", "CONSISTENT", "session_id used everywhere"],
    ],
    col_widths=[2.2, 1.6, 3.0]
)

add_paragraph("14.4 Unused / Orphan Tables (Logical)", bold=True)
add_table(
    ["Orphan Element", "Type", "Referenced By", "Status"],
    [
        ["Student model", "Pydantic model", "Nowhere — not used in repository or engine", "DEFINED BUT UNUSED"],
        ["Device model", "Pydantic model", "DeviceCapabilities, cli doctor", "PARTIALLY USED"],
        ["CurrentActivity", "Pydantic model", "ActivityService", "ISOLATED — never output to API"],
        ["MonitoringSession", "Pydantic model", "StateEngine", "USED INTERNALLY"],
        ["ActionStatus.SKIPPED", "Enum", "ActionEngine idempotency", "USED"],
        ["ErrorActionExecutionError", "Exception", "Core errors", "NEVER RAISED in current code"],
    ],
    col_widths=[2.0, 1.4, 2.2, 1.8]
)

add_page_break()

# ============================================================
# 15. DATABASE DEPENDENCY GRAPH
# ============================================================
doc.add_heading("15. DATABASE DEPENDENCY GRAPH", level=1)

add_paragraph("The dependency graph shows which modules depend on the in-memory data store.", bold=True)

add_code_block(
    "                   DEPENDENCY GRAPH (IN-MEMORY DATA STORE)\n"
    "                   =========================================\n"
    "\n"
    "                          +----------------+\n"
    "                          |   server.py    |\n"
    "                          | (FastAPI)      |\n"
    "                          +----------------+\n"
    "                                |\n"
    "                                v\n"
    "                          +----------------+\n"
    "                          |  ActivityEngine |\n"
    "                          |   (facade)     |\n"
    "                          +----------------+\n"
    "                                |\n"
    "              +----------------+----------------+\n"
    "              |                |                |\n"
    "              v                v                v\n"
    "  +------------+    +----------+      +-----------------+\n"
    "  |  Event     |    |  Policy  |      |   State Engine  |\n"
    "  |  Engine    |    |  Engine  |      |                 |\n"
    "  +------------+    +----------+      +-----------------+\n"
    "        |                |                    |\n"
    "        |                v                    v\n"
    "        |        +----------------+   +-----------------+\n"
    "        |        |  Escalation    |   |   Sessions/     |\n"
    "        |        |  Engine        |   |   Snapshots/    |\n"
    "        |        +----------------+   |   Machines      |\n"
    "        |                |            +-----------------+\n"
    "        |                v                    |\n"
    "        |        +----------------+            |\n"
    "        |        |  Action Engine |            |\n"
    "        |        +----------------+            |\n"
    "        |             |                        |\n"
    "        v             v                        v\n"
    "  +--------------------------------------------------+\n"
    "  |         InMemoryActivityRepository               |\n"
    "  |  events[]      states{}       violations[]      |\n"
    "  +--------------------------------------------------+\n"
    "                          ^\n"
    "                          |\n"
    "  +--------------------------------------------------+"
)


add_paragraph("Key dependency notes:", bold=True)
add_bullet("The InMemoryActivityRepository is the closest thing to a database in this project")
add_bullet("All five engines (Event, Policy, State, Escalation, Action) depend on in-memory data")
add_bullet("server.py (FastAPI) and cli/main.py (CLI) are both entry points that instantiate the ActivityEngine facade")
add_bullet("MonitoringService and FocusService are independent but both feed into the same engines")
add_bullet("The frontend (index.html/script.js) does NOT depend on any of these modules — it is standalone")

add_page_break()

# ============================================================
# 16. CRITICAL COMPONENTS
# ============================================================
doc.add_heading("16. CRITICAL COMPONENTS", level=1)

add_paragraph("The most important architectural components that everything else depends on:", bold=True)

add_table(
    ["Component", "File", "Role", "Criticality"],
    [
        ["ActivityEngine (Facade)", "engine/facade.py", "Single entry point wiring all engines + repository", "CRITICAL"],
        ["EventEngine", "engine/event_engine.py", "Normalizes/deduplicates/staleness-checks all telemetry", "CRITICAL"],
        ["PolicyEngine", "engine/policy_engine.py", "Evaluates policy, escalates, tracks violations", "CRITICAL"],
        ["StateEngine", "engine/state_engine.py", "Maintains state machine, snapshots, sessions", "CRITICAL"],
        ["InMemoryActivityRepository", "persistence/repository.py", "Core data storage for events/states/violations", "CRITICAL"],
        ["PolicyEvaluator", "policy/evaluator.py", "Pure policy logic — no side effects", "HIGH"],
        ["ActionEngine", "engine/action_engine.py", "Executes enforcement actions with idempotency/mode", "HIGH"],
        ["EscalationEngine", "engine/escalation_engine.py", "Plans actions from decisions", "HIGH"],
        ["MonitoringService", "services/monitoring_service.py", "Polls adapters and feeds raw telemetry", "HIGH"],
        ["ConfigLoader", "config/loader.py", "Loads YAML config; fail-fast validation", "MEDIUM"],
        ["server.py", "server.py", "FastAPI/WebSocket entry point", "MEDIUM"],
        ["CLI main", "cli/main.py", "CLI commands for doctor/monitor/simulate/policy-check", "MEDIUM"],
    ],
    col_widths=[1.7, 1.5, 3.0, 1.2]
)

add_paragraph("16.1 What Happens if Each Component Fails", bold=True)
add_table(
    ["Component Fails", "Impact", "Recovery"],
    [
        ["InMemoryActivityRepository", "All events/states/violations lost; API returns empty", "None — process restart loses all data"],
        ["EventEngine", "Telemetry cannot be normalized; raw events dropped", "Restart process"],
        ["PolicyEngine", "No policy decisions made; no enforcement", "Restart process"],
        ["StateEngine", "No state tracking; API /api/state returns UNKNOWN", "Restart process"],
        ["ActionEngine", "No enforcement actions executed", "Restart process"],
        ["MonitoringService", "No telemetry polling from monitors", "Restart process"],
        ["ActivityEngine facade", "Entire engine unavailable", "Restart process"],
        ["server.py", "API endpoints unavailable", "Restart uvicorn"],
    ],
    col_widths=[1.8, 3.2, 1.8]
)

add_page_break()

# ============================================================
# 17. CRITICAL RISKS
# ============================================================
doc.add_heading("17. CRITICAL RISKS", level=1)

add_paragraph("Summary of all identified risks sorted by severity.", bold=True)

add_paragraph("CRITICAL RISKS", bold=True)
add_table(
    ["ID", "Risk", "Severity", "Impact"],
    [
        ["C1", "No persistent database — all data lost on process restart", "CRITICAL", "Complete loss of monitoring data, violations history, state"],
        ["C2", "No backup/recovery mechanism", "CRITICAL", "No way to recover data after crash/restart"],
        ["C3", "No audit trail beyond 1000 most recent events", "CRITICAL", "Historical activity data inaccessible; cannot verify policy compliance over time"],
    ],
    col_widths=[0.5, 3.8, 1.0, 2.0]
)

add_paragraph("HIGH RISKS", bold=True)
add_table(
    ["ID", "Risk", "Severity", "Impact"],
    [
        ["H1", "No authentication on FastAPI server", "HIGH", "Anyone can start/stop sessions, inject telemetry, read state"],
        ["H2", "Gemini API key stored in localStorage / .env.local", "HIGH", "Key theft via XSS or browser access"],
        ["H3", "Violation tracking duplicated (PolicyEngine + Repository) with no sync", "HIGH", "Violation data inconsistency; cannot rely on repository violations"],
        ["H4", "Multiple unbounded in-memory dictionaries (ActionEngine._completed, StateEngine._sessions)", "HIGH", "Memory leaks in long-running processes"],
    ],
    col_widths=[0.5, 3.8, 1.0, 2.0]
)

add_paragraph("MEDIUM RISKS", bold=True)
add_table(
    ["ID", "Risk", "Severity", "Impact"],
    [
        ["M1", "No transaction boundary in facade._route()", "MEDIUM", "Event stored before state — inconsistent snapshot possible"],
        ["M2", "No pagination on API reads", "MEDIUM", "Cannot page through event history"],
        ["M3", "Full event/violation history lost after 1000", "MEDIUM", "No long-term audit trail"],
        ["M4", ".env.local served as static file", "MEDIUM", "Potential API key exposure"],
        ["M5", "EventEngine.process() notifies subscribers on duplicates", "MEDIUM", "Inconsistent dedup behavior between process() and process_raw()"],
        ["M6", "CORS allows all methods/headers", "MEDIUM", "Broad cross-origin access to API"],
    ],
    col_widths=[0.5, 3.8, 1.0, 2.0]
)

add_paragraph("LOW RISKS", bold=True)
add_table(
    ["ID", "Risk", "Severity", "Impact"],
    [
        ["L1", "get_violations() O(n) linear scan", "LOW", "Minor performance impact at 1000 items"],
        ["L2", "get_events() copies list on every read", "LOW", "Minor memory overhead"],
        ["L3", "BedtimeService uses type() instead of WarningRequest class", "LOW", "Type safety violation; future bugs possible"],
        ["L4", "Device IDs inconsistent across code", "LOW", "Confusion in logs/debugging"],
        ["L5", "Several ports/models are dead code", "LOW", "Code maintenance overhead"],
    ],
    col_widths=[0.5, 3.8, 1.0, 2.0]
)

add_page_break()

# ============================================================
# 18. RECOMMENDATIONS
# ============================================================
doc.add_heading("18. RECOMMENDATIONS", level=1)

add_paragraph("Prioritized recommendations to address the risks identified above.", bold=True)

add_paragraph("18.1 Immediate (Critical)", bold=True)
add_table(
    ["#", "Recommendation", "Addresses"],
    [
        ["1", "Implement a persistent repository (SQLite first, PostgreSQL later) implementing the existing ActivityRepository protocol", "C1, C2, C3"],
        ["2", "Add automatic backup/snapshot for in-memory data (periodic flush to disk/DB)", "C2"],
        ["3", "Increase or remove the 1000-event ring buffer; implement pagination with cursors", "C3, M2, M3"],
    ],
    col_widths=[0.4, 5.0, 1.2]
)

add_paragraph("18.2 Short-term (High)", bold=True)
add_table(
    ["#", "Recommendation", "Addresses"],
    [
        ["4", "Add authentication (API key/bearer token) to FastAPI server", "H1"],
        ["5", "Move AI API keys to server-side; never store in localStorage/.env.local", "H2, M4"],
        ["6", "Eliminate duplicate violation storage — route all violations through repository.store_violation()", "H3"],
        ["7", "Add pruning for ActionEngine._completed and StateEngine._sessions", "H4"],
        ["8", "Restrict CORS to specific origins/methods", "M6"],
    ],
    col_widths=[0.4, 5.0, 1.2]
)

add_paragraph("18.3 Medium-term (Medium)", bold=True)
add_table(
    ["#", "Recommendation", "Addresses"],
    [
        ["9", "Wrap facade._route() in transaction-like semantics (store event + state together)", "M1"],
        ["10", "Fix EventEngine.process() to match process_raw() dedup behavior", "M5"],
        ["11", "Add proper data retention and archival for event history", "M3"],
        ["12", "Implement proper pagination with limit/offset or cursor-based", "M2"],
    ],
    col_widths=[0.4, 5.0, 1.2]
)

add_paragraph("18.4 Long-term (Low)", bold=True)
add_table(
    ["#", "Recommendation", "Addresses"],
    [
        ["13", "Replace type() in BedtimeService with real WarningRequest class", "L3"],
        ["14", "Standardize device ID generation across all modules", "L4"],
        ["15", "Remove or implement dead code (ports, unused models)", "L5"],
        ["16", "Add proper data layer with migrations when real DB is added", "C1, C2"],
    ],
    col_widths=[0.4, 5.0, 1.2]
)

add_page_break()

# ============================================================
# 19. HOW THE SYSTEM WORKS — FOR NEW DEVELOPERS
# ============================================================
doc.add_heading("19. HOW THE SYSTEM WORKS — FOR NEW DEVELOPERS", level=1)

add_paragraph(
    "This section is written as onboarding documentation for a developer who is completely new to "
    "the Flowsink project. Read this first before diving into code."
)

add_paragraph("19.1 What is Flowsink?", bold=True)
add_paragraph(
    "Flowsink (repo name) contains two distinct components:\n\n"
    "1. **FocusMate (Frontend)** — A standalone HTML/CSS/JS web app that lets students plan their "
    "weekly schedule (time slots x days), then sends that schedule to Google Gemini or an "
    "OpenAI-compatible model for medical/circadian rhythm analysis and optimization recommendations.\n\n"
    "2. **Student Activity Engine (Python backend)** — A Clean Architecture / Hexagonal / Event-Driven "
    "monitoring engine that receives telemetry (process focus, browser tabs, DNS queries) from "
    "monitor adapters, normalizes it into canonical events, evaluates them against configurable "
    "policies, tracks state transitions, and can enforce actions (warnings, redirects, domain blocks, restricted mode)."
)

add_paragraph("19.2 What database does it use?", bold=True)
add_paragraph(
    "NO DATABASE. Phase 1 uses an in-memory repository (InMemoryActivityRepository) that stores "
    "everything in Python lists/dicts in RAM. Data is lost when the process exits. "
    "The repository interface is protocol-based, making it clear how to add SQLite/PostgreSQL/Redis later."
)

add_paragraph("19.3 How do the components connect?", bold=True)
add_table(
    ["Component", "Location", "Role", "Entry Points"],
    [
        ["FastAPI Server", "src/activity_engine/server.py", "REST + WebSocket API for remote management", "uvicorn src.activity_engine.server:app"],
        ["CLI", "src/activity_engine/cli/main.py", "Doctor, monitor, simulate, policy-check commands", "activity-engine monitor"],
        ["ActivityEngine", "src/activity_engine/engine/facade.py", "Wires all engines + repository together", "ActivityEngine(config, policy, ...)"],
        ["EventEngine", "src/activity_engine/engine/event_engine.py", "Normalize/dedupe/staleness", "engine.feed_raw(dict)"],
        ["PolicyEngine", "src/activity_engine/engine/policy_engine.py", "Evaluate events against policy", "policy_engine.evaluate(event)"],
        ["StateEngine", "src/activity_engine/engine/state_engine.py", "State machine + snapshots", "state_engine.apply_decision(decision)"],
        ["EscalationEngine", "src/activity_engine/engine/escalation_engine.py", "Plan actions", "escalation_engine.plan(decision)"],
        ["ActionEngine", "src/activity_engine/engine/action_engine.py", "Execute actions", "action_engine.execute(request)"],
        ["Repository", "src/activity_engine/persistence/repository.py", "In-memory data store", "InMemoryActivityRepository()"],
        ["Config", "src/activity_engine/config/", "YAML config", "ConfigLoader(path).load()"],
        ["Policy", "src/activity_engine/policy/", "Policy YAML loading", "PolicyLoader(path).load()"],
    ],
    col_widths=[1.2, 1.8, 2.2, 1.8]
)

add_paragraph("19.4 End-to-end flow example", bold=True)
add_code_block(
    "EXAMPLE: Student visits youtube.com\n"
    "================================\n"
    "\n"
    "1. Browser monitor adapter detects new tab\n"
    "2. MonitoringService polls adapter -> gets raw dict:\n"
    "   {\"kind\": \"browser_navigation\", \"browser\": {\"domain\": \"youtube.com\", \"tab_id\": \"tab-1\"}}\n"
    "3. MonitoringService calls event_engine.process_raw(raw)\n"
    "4. EventEngine normalizes -> ActivityEvent(source=browser, type=WEB_NAVIGATION,\n"
    "   browser.domain=youtube.com)\n"
    "5. EventEngine notifies subscribers -> facade._route()\n"
    "6. PolicyEngine.evaluate(event)\n"
    "   -> browser_blocked_domain -> WARNING, level_1, actions=[WARN]\n"
    "7. StateEngine.apply_decision(decision)\n"
    "   -> ActivityStateMachine.transition -> state=WARNING\n"
    "8. EscalationEngine.plan(decision)\n"
    "   -> [ActionRequest(WARN, target=youtube.com)]\n"
    "9. ActionEngine.execute(ActionRequest)\n"
    "   -> if mode=enforce -> adapter.warn() -> SUCCESS\n"
    "   -> if mode=dry_run -> NOT_EXECUTED\n"
    "10. Repository.store_event(event) + Repository.store_state(snapshot)\n"
    "11. API GET /api/state returns WARNING state\n"
)

add_paragraph("19.5 Data model quick reference", bold=True)
add_table(
    ["Model", "Fields (key)", "Stored Where"],
    [
        ["ActivityEvent", "event_id, device_id, source, type, timestamp", "repository._events"],
        ["StateSnapshot", "student_id, device_id, state, risk_score", "repository._states"],
        ["ViolationRecord", "violation_id, student_id, level, risk_score", "policy_engine._violations (NOT in repository)"],
        ["MonitoringSession", "session_id, student_id, device_id, started_at", "state_engine._sessions"],
        ["PolicyDecision", "event_id, outcome, level, action_types, risk_score", "Transient (not stored)"],
        ["ActionRequest", "action_id, action, target, student_id", "Transient (not stored)"],
        ["ActionResult", "action_id, action, status, timestamp", "Transient (not stored)"],
    ],
    col_widths=[1.4, 3.0, 2.8]
)

add_paragraph("19.6 Testing the engine", bold=True)
add_bullet("Run pytest: 'pytest' (requires dev dependencies)")
add_bullet("Run CLI simulation: 'activity-engine simulate --events 10 --mode dry_run'")
add_bullet("Run CLI monitor: 'activity-engine monitor --poll 3 --mode dry_run'")
add_bullet("Start server: 'uvicorn src.activity_engine.server:app --reload --port 8000'")
add_bullet("Use MockActionExecutor for unit tests (adapters/mock/action_executor.py)")

add_page_break()

# ============================================================
# 20. APPENDIX
# ============================================================
doc.add_heading("20. APPENDIX", level=1)

add_paragraph("20.1 Files Scanned During Audit", bold=True)
add_table(
    ["File", "Status"],
    [
        ["src/activity_engine/__init__.py", "Scanned"],
        ["src/activity_engine/server.py", "Scanned"],
        ["src/activity_engine/engine/facade.py", "Scanned"],
        ["src/activity_engine/engine/event_engine.py", "Scanned"],
        ["src/activity_engine/engine/policy_engine.py", "Scanned"],
        ["src/activity_engine/engine/state_engine.py", "Scanned"],
        ["src/activity_engine/engine/escalation_engine.py", "Scanned"],
        ["src/activity_engine/engine/action_engine.py", "Scanned"],
        ["src/activity_engine/persistence/repository.py", "Scanned"],
        ["src/activity_engine/policy/evaluator.py", "Scanned"],
        ["src/activity_engine/policy/classifier.py", "Scanned"],
        ["src/activity_engine/policy/loader.py", "Scanned"],
        ["src/activity_engine/policy/default_policies.yaml", "Scanned"],
        ["src/activity_engine/services/activity_service.py", "Scanned"],
        ["src/activity_engine/services/bedtime_service.py", "Scanned"],
        ["src/activity_engine/services/focus_service.py", "Scanned"],
        ["src/activity_engine/services/monitoring_service.py", "Scanned"],
        ["src/activity_engine/core/events.py", "Scanned"],
        ["src/activity_engine/core/states.py", "Scanned"],
        ["src/activity_engine/core/actions.py", "Scanned"],
        ["src/activity_engine/core/decisions.py", "Scanned"],
        ["src/activity_engine/core/policies.py", "Scanned"],
        ["src/activity_engine/core/errors.py", "Scanned"],
        ["src/activity_engine/domain/session.py", "Scanned"],
        ["src/activity_engine/domain/student.py", "Scanned"],
        ["src/activity_engine/domain/device.py", "Scanned"],
        ["src/activity_engine/domain/activity.py", "Scanned"],
        ["src/activity_engine/domain/violations.py", "Scanned"],
        ["src/activity_engine/ports/__init__.py", "Scanned"],
        ["src/activity_engine/adapters/mock/*.py", "Scanned"],
        ["src/activity_engine/transport/dto.py", "Scanned"],
        ["src/activity_engine/transport/protocol.py", "Scanned"],
        ["src/activity_engine/transport/serialization.py", "Scanned"],
        ["src/activity_engine/config/models.py", "Scanned"],
        ["src/activity_engine/config/loader.py", "Scanned"],
        ["src/activity_engine/cli/main.py", "Scanned"],
        ["tests/conftest.py", "Scanned"],
        ["tests/test_integration.py", "Scanned"],
        ["tests/test_policy_engine.py", "Scanned"],
        ["tests/test_states.py", "Scanned"],
        ["tests/test_contract.py", "Scanned"],
        ["examples/simulate.py", "Scanned"],
        ["examples/policy.yaml", "Scanned"],
        ["docs/ARCHITECTURE.md", "Scanned"],
        ["docs/INTEGRATION.md", "Scanned"],
        ["docs/EVENT_CONTRACT.md", "Scanned"],
        ["docs/POLICY_ENGINE.md", "Scanned"],
        ["docs/SECURITY.md", "Scanned"],
        ["pyproject.toml", "Scanned"],
        ["package.json", "Scanned (no dependencies — not an npm project)"],
        [".gitignore", "Scanned"],
        ["README.md", "Scanned"],
        ["index.html", "Scanned"],
        ["script.js", "Scanned"],
        ["style.css", "Scanned"],
        ["config/ (directory)", "NOT PRESENT — config loaded from config/policy.yaml path but dir absent"],
    ],
    col_widths=[3.5, 2.5]
)

add_paragraph("20.2 Database Indicators — Complete Negative Results", bold=True)
add_table(
    ["Indicator", "Result"],
    [
        ["Database engines", "NONE found"],
        ["Database files", "NONE found (*.db, *.sqlite, *.sqlite3)"],
        ["SQL files", "NONE found (*.sql)"],
        ["ORM files", "NONE found (SQLAlchemy, Django ORM, Tortoise, Peewee)"],
        ["Migration files", "NONE found (alembic, django migrations)"],
        ["Connection strings", "NONE found"],
        ["DATABASE_URL", "NONE found"],
        [".env files", "NONE found"],
        ["Docker files", "NONE found (Dockerfile, docker-compose.yml)"],
        ["Seeds/scripts", "NONE found"],
        ["Redis/Supabase/Firebase", "NONE found"],
    ],
    col_widths=[2.2, 3.8]
)

add_paragraph("20.3 Methodology", bold=True)
add_bullet("Phase 1: Recursively scanned all project directories for database indicators")
add_bullet("Phase 2: Read complete contents of all source files (core, engine, services, adapters, ports, transport, config, cli)")
add_bullet("Phase 3: Read test suite (conftest, integration, policy, states, contract tests)")
add_bullet("Phase 4: Read documentation (architecture, integration, event contract, policy engine, security)")
add_bullet("Phase 5: Read frontend files (index.html, script.js, style.css)")
add_bullet("Phase 6: Cross-referenced all imports to identify storage dependents")
add_bullet("Phase 7: Built complete schema, data flow, query flow, and dependency maps")
add_bullet("Phase 8: Authored this DOCX report with python-docx")

add_paragraph("20.4 Audit Scope Limitations", bold=True)
add_bullet("This audit covers source code present in the workspace only")
add_bullet("Runtime state (memory contents) could not be inspected — no running processes")
add_bullet("Browser localStorage contents not inspected")
add_bullet("No network calls made; no external system states inspected")

# ============================================================
# FINAL SAVE
# ============================================================
add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("END OF REPORT")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUTPUT_FILE)
print(f"Report generated successfully: {OUTPUT_FILE}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
